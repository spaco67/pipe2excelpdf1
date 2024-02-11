from flask import Flask, request, render_template, send_from_directory, redirect, url_for
import pandas as pd
import os
from datetime import datetime
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = 'temp_files'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        data = request.form['data']
        if data:
            try:
                df = pd.DataFrame([x.split('|') for x in data.split('\n')])
                new_header = df.iloc[0]  # First line as header
                df = df[1:]
                df.columns = new_header

                timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
                if 'convert_excel' in request.form:
                    filename = secure_filename(f'converted_{timestamp}.xlsx')
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    df.to_excel(filepath, index=False)
                elif 'convert_pdf' in request.form:
                    filename = secure_filename(f'converted_{timestamp}.pdf')
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    doc = SimpleDocTemplate(filepath, pagesize=letter)
                    data = [df.columns.to_list()] + df.values.tolist()
                    table = Table(data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    doc.build([table])
                elif 'convert_docx' in request.form:
                    filename = secure_filename(f'converted_{timestamp}.docx')
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    doc = Document()
                    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                    for j, col in enumerate(df.columns):
                        table.cell(0, j).text = col
                        table.cell(0, j).paragraphs[0].runs[0].font.bold = True
                    for i, row in enumerate(df.values):
                        for j, val in enumerate(row):
                            table.cell(i+1, j).text = str(val)
                    doc.save(filepath)

                return redirect(url_for('download_file', filename=filename))

            except Exception as e:
                return render_template('index.html', error=f"An error occurred: {e}")

    return render_template('index.html', error=None)

@app.route('/downloads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
